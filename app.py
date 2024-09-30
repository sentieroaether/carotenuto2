import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
import zipfile
from docx2pdf import convert  # Per convertire il documento Word in PDF
import os
import pythoncom  # Assicurati di importarlo correttamente
from docx2pdf import convert


# Funzione per caricare i file Excel o CSV
def carica_file():
    uploaded_file1 = st.file_uploader("Carica il file Anagrafiche (Excel o CSV)", type=["xlsx", "csv"])
    uploaded_file2 = st.file_uploader("Carica il file Fatture (Excel o CSV)", type=["xlsx", "csv"])
    uploaded_file3 = st.file_uploader("Carica il file Pratiche (Excel o CSV)", type=["xlsx", "csv"])
    
    def leggi_file(file):
        if file is not None:
            try:
                if file.name.endswith('.xlsx'):
                    return pd.read_excel(file)
                elif file.name.endswith('.csv'):
                    return pd.read_csv(file, sep=';', on_bad_lines='skip')
            except Exception as e:
                st.error(f"Errore nel caricamento del file {file.name}: {e}")
        return None

    df1 = leggi_file(uploaded_file1)  # File Anagrafiche
    df2 = leggi_file(uploaded_file2)  # File Fatture
    df3 = leggi_file(uploaded_file3)  # File Pratiche
    
    if df1 is None or df2 is None or df3 is None:
        st.warning("Uno o più file non sono stati caricati correttamente.")
        return None, None, None
    
    return df1, df2, df3

# Funzione per normalizzare i nomi delle colonne (convertirli in minuscolo)
def normalizza_colonne(df):
    df.columns = df.columns.str.lower()  # Converte tutte le colonne in minuscolo
    return df

# Funzione per applicare il grassetto e la dimensione 12 al testo
def set_bold_and_size(paragraph, bold_text):
    for run in paragraph.runs:
        if bold_text in run.text:
            run.bold = True
        run.font.size = Pt(12)  # Imposta la dimensione del font a 12

# Funzione per generare il documento Word con sostituzioni e formattazione
def genera_documento_word(dati, template_path="decreto.docx"):
    st.write("Dati per la sostituzione:", dati)
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        # Sostituzione dei segnaposto con i valori forniti o con stringa vuota se il valore è 'NaN'
        placeholders = {
            "{ragione_sociale}": str(dati.get('ragione_sociale', '')).replace('nan', ''),  # Nome corretto
            "{codice_fiscale}": str(dati.get('codice_fiscale', '')).replace('nan', ''),
            "{partita_va}": str(dati.get('partita_iva', '')).replace('nan', ''),
            "{comune_residenza}": str(dati.get('comune_residenza', '')).replace('nan', ''),
            "{cap_residenza}": str(int(dati.get('cap_residenza', 0))),  # Rimuove i decimali dal CAP
            "{indirizzo_residenza}": str(dati.get('indirizzo_residenza', '')).replace('nan', ''),
            "{settore_contabile}": str(dati.get('settore_contabile', '')).replace('nan', ''),
            "{codice_commerciale}": str(dati.get('codice_commerciale', '')).replace('nan', ''),
            "{codice_soggetto}": str(int(dati.get('codice_soggetto', 0))),  # Rimuove i decimali .0
            "{comune_fornitura}": str(dati.get('comune_fornitura', '')).replace('nan', ''),
            "{provincia_fornitura}": str(dati.get('provincia_fornitura', '')).replace('nan', ''),
            "{indirizzo_fornitura}": str(dati.get('indirizzo_fornitura', '')).replace('nan', ''),
            "{pod}": str(dati.get('pod', '')).replace('nan', ''),
            "{residuo_ad_oggi}": str(dati.get('residuo ad oggi', '')).replace('nan', ''),
        }
        
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                set_bold_and_size(paragraph, placeholder)  # Imposta il grassetto per il campo

        # Imposta il grassetto per "Ragione Sociale" e "Residuo ad oggi"
        if "ragione_sociale" in paragraph.text or "residuo_ad_oggi" in paragraph.text:
            for run in paragraph.runs:
                run.bold = True

    # Imposta la dimensione del font a 12 per tutti i paragrafi
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Funzione per convertire un documento Word in PDF
def convert_to_pdf(word_file, output_pdf_path):
    pythoncom.CoInitialize()  # Inizializza COM

    with open("temp_decreto.docx", "wb") as temp_word_file:
        temp_word_file.write(word_file.getvalue())
    
    convert("temp_decreto.docx", output_pdf_path)
    
    # Rimuovi il file temporaneo dopo la conversione
    os.remove("temp_decreto.docx")
    
# Funzione principale
def main():
    st.title("Generatore di Documenti Word e PDF da Excel o CSV")
    df1, df2, df3 = carica_file()
    
    if df1 is not None and df2 is not None and df3 is not None:
        st.write("File caricati con successo!")
        
        try:
            # Normalizza i nomi delle colonne a minuscolo
            df1 = normalizza_colonne(df1)
            df2 = normalizza_colonne(df2)
            df3 = normalizza_colonne(df3)

            # Rinominazione delle colonne per fare il merge
            df1 = df1.rename(columns={"codice_soggetto": "codice_soggetto", "ragione_sociale": "ragione_sociale", "numero affido_x": "numero affido_x"})  # Usa Numero affido_x
            df2 = df2.rename(columns={"bpartner": "codice_soggetto"})
            df3 = df3.rename(columns={"soggetto": "codice_soggetto"})
        except KeyError as e:
            st.error(f"Errore durante la rinominazione delle colonne: {e}")
            return
        
        if "codice_soggetto" not in df1.columns or "codice_soggetto" not in df2.columns or "codice_soggetto" not in df3.columns:
            st.error("Una o più colonne 'codice_soggetto' non sono state trovate.")
            return
        
        try:
            df_combinato = pd.merge(df1, df2, on='codice_soggetto')
            df_combinato = pd.merge(df_combinato, df3, on='codice_soggetto')
        except KeyError as e:
            st.error(f"Errore durante l'unione dei file: {e}")
            return
        
        codici_soggetto = df_combinato['codice_soggetto'].unique()
        codici_selezionati = st.multiselect("Seleziona i codici soggetto per generare i documenti:", codici_soggetto)
        
        if codici_selezionati:
            zip_buffer = BytesIO()
            
            with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                for codice in codici_selezionati:
                    dati_filtro = df_combinato[df_combinato['codice_soggetto'] == codice].iloc[0]
                    doc_buffer = genera_documento_word(dati_filtro, template_path="decreto.docx")
                    
                    # Nome del file personalizzato: 'Ragione_Sociale' + 'Codice_Soggetto' + 'Numero_Affido_x'
                    nome_file = f"{dati_filtro['ragione_sociale']}_{int(dati_filtro['codice_soggetto'])}_{dati_filtro['numero affido_x']}"
                    
                    # Aggiungi il documento Word allo zip
                    zip_file.writestr(f"{nome_file}.docx", doc_buffer.getvalue())
                    
                    # Genera anche il PDF e aggiungilo allo zip
                    pdf_path = f"{nome_file}.pdf"
                    convert_to_pdf(doc_buffer, pdf_path)
                    with open(pdf_path, "rb") as pdf_file:
                        zip_file.writestr(f"{nome_file}.pdf", pdf_file.read())
                    
                    # Rimuovi il PDF generato
                    os.remove(pdf_path)
            
            zip_buffer.seek(0)
            st.download_button(
                label="Scarica tutti i documenti generati (ZIP)",
                data=zip_buffer,
                file_name="documenti_generati.zip",
                mime="application/zip"
            )
    else:
         st.warning("Per favore, carica tutti e tre i file (Excel o CSV).")

if __name__ == "__main__":
    main()
