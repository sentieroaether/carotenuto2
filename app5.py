import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import os
import pythoncom
import win32com.client as win32
from datetime import datetime
import numpy as np


def rimuovi_decimali(valore):
    try:
        return str(int(float(valore)))
    except (ValueError, TypeError):
        return str(valore)

st.markdown("[Clicca qui per collegarti al Portale Automazione Lettere Diffida](https://pagp-slcarotenuto.streamlit.app/)", unsafe_allow_html=True)
# Funzione per formattare i numeri (POD, codice soggetto) come interi
def formatta_numero_intero(numero):
    try:
        return str(int(float(numero)))
    except ValueError:
        return str(numero)

def formatta_data_italiana(data):
    mesi = [
        "gennaio", "febbraio", "marzo", "aprile", "maggio", "giugno",
        "luglio", "agosto", "settembre", "ottobre", "novembre", "dicembre"
    ]
    giorno = data.day
    mese = mesi[data.month - 1]
    anno = data.year
    return f"{giorno} {mese} {anno}"

# Funzione per convertire un documento Word in PDF
def convert_to_pdf(word_file, output_pdf_path):
    # Inizializza COM
    pythoncom.CoInitialize()

    try:
        with open("temp_decreto.docx", "wb") as temp_word_file:
            temp_word_file.write(word_file.getvalue())

        # Crea un'istanza di Word
        word = win32.Dispatch("Word.Application")
        word.Visible = False

        # Apri il documento temporaneo
        doc = word.Documents.Open(os.path.abspath("temp_decreto.docx"))

        # Salva il documento come PDF
        doc.SaveAs(os.path.abspath(output_pdf_path), FileFormat=17)  # 17 è il formato PDF in Word
        doc.Close()

        # Chiudi l'applicazione Word
        word.Quit()

        # Rimuovi il file temporaneo
        if os.path.exists("temp_decreto.docx"):
            os.remove("temp_decreto.docx")
    
    finally:
        # Deinizializza COM
        pythoncom.CoUninitialize()

# Funzione per caricare i file Excel o CSV
def carica_file():
    uploaded_file1 = st.file_uploader("Carica il file ANAGRAFICHE", type=["xlsx", "csv"])
    uploaded_file2 = st.file_uploader("Carica il file FATTURE", type=["xlsx", "csv"])
    uploaded_file3 = st.file_uploader("Carica il file PRATICHE ", type=["xlsx", "csv"])
    
    def leggi_file(file):
        if file is not None:
            try:
                if file.name.endswith('.xlsx'):
                    return pd.read_excel(file)
                elif file.name.endswith('.csv'):
                    return pd.read_csv(file, sep=';', on_bad_lines='skip')
            except Exception as e:
                st.error(f"Errore nel caricamento del file {file.name}: {e}")
        return None

    df1 = leggi_file(uploaded_file1)  # File Anagrafiche
    df2 = leggi_file(uploaded_file2)  # File Fatture
    df3 = leggi_file(uploaded_file3)  # File Pratiche
    
    return df1, df2, df3

# Funzione per normalizzare i nomi delle colonne (convertirli in minuscolo e rimuovere caratteri speciali)
def normalizza_colonne(df):
    df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_').str.replace('.', '').str.replace("'", "")
    return df

def compila_tabella_esistente(doc, df_combinato):
    try:
        # Filtra le righe che hanno un NaN in n_documento
        df_combinato = df_combinato.dropna(subset=['n_documento'])
        
        # Prendi la prima tabella del documento
        table = doc.tables[0]

        # Numero di righe di dati da inserire
        righe_da_aggiungere = len(df_combinato)

        # Numero di righe esistenti nella tabella Word
        righe_esistenti = len(table.rows)
        
        # Aggiungi righe necessarie per ciascuna fattura se la tabella ha meno righe di quante servono
        if righe_da_aggiungere > righe_esistenti - 1:  # Tolgo 1 per la riga di intestazione
            for _ in range(righe_da_aggiungere - (righe_esistenti - 1)):
                table.add_row()

        # Popola la tabella con le fatture
        for i, row in enumerate(df_combinato.itertuples(), 1):  # Partiamo dalla seconda riga, indice 1 (la prima è l'intestazione)
            cells = table.rows[i].cells  # Accedi alla riga corretta della tabella

            # Popola ogni cella con i dati corretti
            cells[0].text = str(row.data_reg)  
            cells[1].text = str(row.scadnetto)
            cells[2].text = str(row.n_documento)
            cells[3].text = str(row.importo_totale)
            cells[4].text = str(row.importo_pagato_totale)
            cells[5].text = str(row.residuo_ad_oggi)
            
            # Gestione del campo POD (sia stringa che float)
            pod_value = row.pod
            cells[6].text = formatta_numero_intero(pod_value)

    except Exception as e:
        st.error(f"Errore durante la compilazione della tabella: {e}")

# Funzione per generare il documento Word con debug per ragione sociale
def genera_documento_word(dati, df_combinato, template_path="decreto.docx"):
    doc = Document(template_path)


    
    data_generazione = formatta_data_italiana(datetime.now())

    # Sostituzione dei segnaposto nel documento
    placeholders = {
    "{ragione_sociale}": str(dati.get('ragione_sociale_x', '')).replace('nan', ''),
    "{codice_fiscale}": rimuovi_decimali(dati.get('codice_fiscale', '')),
    "{partita_iva}": rimuovi_decimali(dati.get('partita_iva', '')),
    "{comune_residenza}": str(dati.get('comune_residenza', '')).replace('nan', ''),
    "{cap_residenza}": rimuovi_decimali(dati.get('cap_residenza', 0)),
    "{indirizzo_residenza}": str(dati.get('indirizzo_residenza', '')).replace('nan', ''),
    "{settore_contabile}": str(dati.get('settore_contabile', '')).replace('nan', ''),
    "{codice_commerciale}": rimuovi_decimali(dati.get('codice_commerciale', '')),
    "{codice_soggetto}": rimuovi_decimali(dati.get('codice_soggetto', 0)),
    "{comune_fornitura}": str(dati.get('comune_fornitura', '')).replace('nan', ''),
    "{provincia_fornitura}": str(dati.get('provincia_fornitura', '')).replace('nan', ''),
    "{indirizzo_fornitura}": str(dati.get('indirizzo_fornitura', '')).replace('nan', ''),
    "{data_generazione}": data_generazione,
    "{provincia_residenza}": str(dati.get('provincia_residenza', '')).replace('nan', ''),
    "{pod}": rimuovi_decimali(dati.get('pod', '')),
    "{residuo_ad_oggi}": rimuovi_decimali(dati.get('residuo_ad_oggi', '')),
}


    # Sostituzione dei segnaposto nel testo del documento
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Compila la tabella esistente nel template
    compila_tabella_esistente(doc, df_combinato)

    # Salva il documento in un buffer BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Funzione principale
def main():
    st.title("Generatore automatico di Ricorsi D.I.")
    df1, df2, df3 = carica_file()
    
    if df1 is not None and df2 is not None and df3 is not None:
        st.write("File caricati con successo!")

        try:
            # Normalizza i nomi delle colonne a minuscolo
            df1 = normalizza_colonne(df1)
            df2 = normalizza_colonne(df2)
            df3 = normalizza_colonne(df3)
 
         
            # Rinominazione delle colonne per fare il merge
            df1 = df1.rename(columns={"codice_soggetto": "codice_soggetto"})
            df2 = df2.rename(columns={"bpartner": "codice_soggetto"})
            df3 = df3.rename(columns={"soggetto": "codice_soggetto"})
        except KeyError as e:
            st.error(f"Errore durante la rinominazione delle colonne: {e}")
            return

        if "codice_soggetto" not in df1.columns or "codice_soggetto" not in df2.columns or "codice_soggetto" not in df3.columns:
            st.error("Una o più colonne 'codice_soggetto' non sono state trovate.")
            return

        try:
            df_combinato = pd.merge(df1, df2, on='codice_soggetto', how='left')  # Unione tra df1 e df2
            df_combinato = pd.merge(df_combinato, df3, on='codice_soggetto', how='left')  # Unione con df3
        except KeyError as e:
            st.error(f"Errore durante l'unione dei file: {e}")
            return

        # Gestione dei codici soggetto senza NaN
        codici_soggetto = df_combinato['codice_soggetto'].unique()
        codici_soggetto = [int(float(c)) for c in codici_soggetto if pd.notna(c)]  # Ignora i NaN

        # Mostra i codici soggetto correttamente formattati nella multiselect
        codici_selezionati = st.multiselect("Seleziona i codici soggetto per generare i documenti:", codici_soggetto)

        if codici_selezionati:
            zip_buffer = BytesIO()

            with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                for codice in codici_selezionati:
                    # Filtra i dati per il codice soggetto selezionato
                    dati_filtro = df_combinato[df_combinato['codice_soggetto'] == codice].iloc[0]

                    # Filtra le fatture corrispondenti a questo codice soggetto
                    df_fatture = df_combinato[df_combinato['codice_soggetto'] == codice]

                    # Genera il documento Word con la tabella delle fatture
                    doc_buffer = genera_documento_word(dati_filtro, df_fatture, template_path="decreto.docx")

                    # Nome del file personalizzato: 'Codice_Soggetto'
                    nome_file = formatta_numero_intero(dati_filtro['codice_soggetto'])

                    # Aggiungi il documento Word allo zip
                    zip_file.writestr(f"{nome_file}.docx", doc_buffer.getvalue())

                    # Genera anche il PDF e aggiungilo allo zip
                    pdf_path = f"{nome_file}.pdf"
                    convert_to_pdf(doc_buffer, pdf_path)
                    with open(pdf_path, "rb") as pdf_file:
                        zip_file.writestr(f"{nome_file}.pdf", pdf_file.read())

                    # Rimuovi il PDF generato
                    os.remove(pdf_path)
            
            zip_buffer.seek(0)
            st.download_button(
                label="Scarica tutti i documenti generati (ZIP)",
                data=zip_buffer,
                file_name="documenti_generati.zip",
                mime="application/zip"
            )
    else:
        st.warning("Per favore, carica tutti e tre i file (Excel o CSV).")

if __name__ == "__main__":
    main()
