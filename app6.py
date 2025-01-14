import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import os
import pythoncom
import win32com.client as win32
from datetime import datetime
import numpy as np
from docx import Document
from io import BytesIO

# Dati di accesso predefiniti
DEFAULT_USERNAME = "admin"
DEFAULT_PASSWORD = "admin"

# Inizializza le variabili di sessione
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if 'username' not in st.session_state:
    st.session_state['username'] = DEFAULT_USERNAME

if 'password' not in st.session_state:
    st.session_state['password'] = DEFAULT_PASSWORD

# Funzione per sostituire NaN con stringa vuota
def valore_o_spazio(valore):
    if pd.isna(valore) or valore == "nan" or valore == None:
        return ""
    return str(valore)

# Funzione per rimuovere i decimali
def rimuovi_decimali(valore):
    try:
        # Applica la rimozione dei decimali solo su valori numerici validi
        return str(int(float(valore)))
    except (ValueError, TypeError):
        # Se il valore non è un numero o non può essere convertito, restituiscilo così com'è
        return str(valore)

def formatta_pod(valore):
    try:
        # Verifica se il valore è numerico e convertilo in intero, mantenendo gli zeri iniziali
        if isinstance(valore, (int, float)):
            # Converte in intero se float e poi in stringa senza perdere zeri iniziali
            valore = '{:0.0f}'.format(float(valore))
        return str(valore).upper().strip()  # Mantiene la formattazione in maiuscolo e rimuove eventuali spazi
    except (ValueError, TypeError):
        # Se c'è un errore, restituisci il valore così com'è
        return str(valore)
    
# Funzione per formattare i numeri come stringa senza notazione scientifica
def formatta_numero_intero(numero):
    try:
        numero_float = float(numero)  # Prima converti in float per evitare errori
        numero_intero = int(numero_float)  # Converti in intero senza notazione scientifica
        return '{:d}'.format(numero_intero)  # Restituisci come stringa senza notazione
    except ValueError:
        return str(numero)  # Se non è un numero valido, restituisci il valore originale come stringa


# Funzione per formattare le date in italiano
def formatta_data_italiana(data):
    mesi = [
        "gennaio", "febbraio", "marzo", "aprile", "maggio", "giugno",
        "luglio", "agosto", "settembre", "ottobre", "novembre", "dicembre"
    ]
    giorno = data.day
    mese = mesi[data.month - 1]
    anno = data.year
    return f"{giorno} {mese} {anno}"

# Funzione per convertire un documento Word in PDF
def convert_to_pdf(word_file, output_pdf_path):
    pythoncom.CoInitialize()

    try:
        # Creazione di un file temporaneo con il nome corretto
        temp_file_path = "temp_decreto.docx"  # File temporaneo
        with open(temp_file_path, "wb") as temp_word_file:
            temp_word_file.write(word_file.getvalue())  # Scrivi il contenuto del documento temporaneo

        word = win32.Dispatch("Word.Application")
        word.Visible = False

        # Apri il file temporaneo Word e convertilo in PDF
        doc = word.Documents.Open(os.path.abspath(temp_file_path))
        doc.SaveAs(os.path.abspath(output_pdf_path), FileFormat=17)  # 17 è il formato PDF in Word
        doc.Close()
        word.Quit()

        # Elimina il file temporaneo dopo la conversione
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

    except Exception as e:
        st.error(f"Errore durante la conversione in PDF: {e}")
    finally:
        pythoncom.CoUninitialize()

# Funzione per caricare i file Excel o CSV
def carica_file():
    uploaded_file1 = st.file_uploader("Carica il file ANAGRAFICHE", type=["xlsx", "csv"])
    uploaded_file2 = st.file_uploader("Carica il file FATTURE", type=["xlsx", "csv"])
    uploaded_file3 = st.file_uploader("Carica il file PRATICHE", type=["xlsx", "csv"])
    
    def leggi_file(file):
        if file is not None:
            try:
                if file.name.endswith('.xlsx'):
                    # Escludi 'NA' dai valori NaN
                    return pd.read_excel(file, na_values=["", "null", "nan", "NaN"], keep_default_na=False)
                elif file.name.endswith('.csv'):
                    # Escludi 'NA' dai valori NaN
                    return pd.read_csv(file, sep=';', na_values=["", "null", "nan", "NaN"], keep_default_na=False)
            except Exception as e:
                st.error(f"Errore nel caricamento del file {file.name}: {e}")
        return None


    df1 = leggi_file(uploaded_file1)
    df2 = leggi_file(uploaded_file2)
    df3 = leggi_file(uploaded_file3)
    
    return df1, df2, df3

# Funzione per normalizzare i nomi delle colonne
def normalizza_colonne(df):
    df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_').str.replace('.', '').str.replace("'", "")
    return df

# Funzione per compilare la tabella nel documento Word
def compila_tabella_esistente(doc, df_combinato):
    try:
        df_combinato = df_combinato.dropna(subset=['n_documento'])
        table = doc.tables[0]
        righe_da_aggiungere = len(df_combinato)
        righe_esistenti = len(table.rows)

        if righe_da_aggiungere > righe_esistenti - 1:
            for _ in range(righe_da_aggiungere - (righe_esistenti - 1)):
                table.add_row()

        for i, row in enumerate(df_combinato.itertuples(), 1):
            cells = table.rows[i].cells
            cells[0].text = str(row.data_reg)  
            cells[1].text = str(row.scadnetto)
            cells[2].text = str(row.n_documento)
            cells[3].text = str(row.importo_totale)
            cells[4].text = str(row.importo_pagato_totale)
            cells[5].text = str(row.residuo_ad_oggi)
            pod_value = row.pod
            cells[6].text = formatta_pod(pod_value)  # Utilizza la funzione per formattare il POD

    except Exception as e:
        st.error(f"Errore durante la compilazione della tabella: {e}")



# Funzione per generare il documento Word
def genera_documento_word(dati, df_combinato, template_path="decreto.docx"):
    doc = Document(template_path)
    data_generazione = formatta_data_italiana(datetime.now())

    placeholders = {
        "{ragione_sociale}": str(valore_o_spazio(dati.get('ragione_sociale_x',''))),
        "{codice_fiscale}": str(rimuovi_decimali(valore_o_spazio(dati.get('codice_fiscale', '')))),
        "{partita_iva}": str(rimuovi_decimali(valore_o_spazio(dati.get('partita_iva', '')))),
        "{comune_residenza}": str(valore_o_spazio(dati.get('comune_residenza', ''))),
        "{cap_residenza}": str(rimuovi_decimali(valore_o_spazio(dati.get('cap_residenza', 0)))),
        "{indirizzo_residenza}": str(valore_o_spazio(dati.get('indirizzo_residenza', ''))),
        "{settore_contabile}": str(valore_o_spazio(dati.get('settore_contabile', ''))),
        "{codice_commerciale}": str(rimuovi_decimali(valore_o_spazio(dati.get('codice_commerciale', '')))),
        "{codice_soggetto}": str(rimuovi_decimali(valore_o_spazio(dati.get('codice_soggetto', 0)))),
        "{comune_fornitura}": str(valore_o_spazio(dati.get('comune_fornitura', ''))),
        "{provincia_fornitura}": str(valore_o_spazio(dati.get('provincia_fornitura', 'NA')).replace('nan', '')),
        "{indirizzo_fornitura}": str(valore_o_spazio(dati.get('indirizzo_fornitura', ''))),
        "{data_generazione}": str(data_generazione),
        "{provincia_residenza}": str(valore_o_spazio(dati.get('provincia_residenza', 'NA')).replace('nan', '')),
        "{pod}": str(formatta_pod(valore_o_spazio(dati.get('pod', '')))),  # Assicurati che anche il POD sia stringa
        "{residuo_ad_oggi}": str(rimuovi_decimali(valore_o_spazio(dati.get('residuo_ad_oggi', '')))),
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

                                   

    compila_tabella_esistente(doc, df_combinato)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Funzione di login
def login():
    st.title("Benvenuti al Portale Ricorsi D.I. dello Studio Legale Associato Carotenuto")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if username == st.session_state['username'] and password == st.session_state['password']:
            st.session_state['authenticated'] = True
            st.success(f"Accesso per Studio Legale Associato Carotenuto!")
        else:
            st.error("Username o password errati!")

# Funzione principale dell'applicazione
def main():
    st.markdown("[Clicca qui per collegarti al Portale Automazione Lettere Diffida](https://pagp-slcarotenuto.streamlit.app/)", unsafe_allow_html=True)

    if st.session_state['authenticated']:
        st.sidebar.title("Portale ricorso D.I. | Studio Legale Associato Carotenuto")

        st.title("Generatore automatico di Ricorsi D.I.")
        df1, df2, df3 = carica_file()
        
        if df1 is not None and df2 is not None and df3 is not None:
            st.write("File caricati con successo!")

            try:
                df1 = normalizza_colonne(df1)
                df2 = normalizza_colonne(df2)
                df3 = normalizza_colonne(df3)

                df1 = df1.rename(columns={"codice_soggetto": "codice_soggetto"})
                df2 = df2.rename(columns={"bpartner": "codice_soggetto"})
                df3 = df3.rename(columns={"soggetto": "codice_soggetto"})
            except KeyError as e:
                st.error(f"Errore durante la rinominazione delle colonne: {e}")
                return

            if "codice_soggetto" not in df1.columns or "codice_soggetto" not in df2.columns or "codice_soggetto" not in df3.columns:
                st.error("Una o più colonne 'codice_soggetto' non sono state trovate.")
                return

            try:
                df_combinato = pd.merge(df1, df2, on='codice_soggetto', how='left')
                df_combinato = pd.merge(df_combinato, df3, on='codice_soggetto', how='left')
                df_combinato['pod'] = df_combinato['pod'].astype(str)
        
            except KeyError as e:
                st.error(f"Errore durante l'unione dei file: {e}")
                return

            codici_soggetto = df_combinato['codice_soggetto'].unique()
            codici_soggetto = [int(float(c)) for c in codici_soggetto if pd.notna(c)]

            codici_selezionati = st.multiselect("Seleziona i codici soggetto per generare i documenti:", codici_soggetto)
            if st.button("Genera documenti per i soggetti selezionati"):
                if codici_selezionati:
                    zip_buffer = BytesIO()

                    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                        for codice in codici_selezionati:
                            dati_filtro = df_combinato[df_combinato['codice_soggetto'] == codice].iloc[0]
                            df_fatture = df_combinato[df_combinato['codice_soggetto'] == codice]
                            doc_buffer = genera_documento_word(dati_filtro, df_fatture, template_path="decreto.docx")

                            nome_file = formatta_numero_intero(dati_filtro['codice_soggetto'])
                            zip_file.writestr(f"{nome_file}.docx", doc_buffer.getvalue())

                            pdf_path = f"{nome_file}.pdf"
                            convert_to_pdf(doc_buffer, pdf_path)
                            with open(pdf_path, "rb") as pdf_file:
                                zip_file.writestr(f"{nome_file}.pdf", pdf_file.read())

                            os.remove(pdf_path)
                    
                    zip_buffer.seek(0)
                    st.download_button(
                        label="Scarica tutti i documenti generati (ZIP)",
                        data=zip_buffer,
                        file_name="documenti_generati.zip",
                        mime="application/zip"
                    )
            else:
                st.warning("Seleziona almeno un codice soggetto prima di generare i documenti.")
    else:
        login()

if __name__ == "__main__":
    main()
