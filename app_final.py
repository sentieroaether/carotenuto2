import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import os
from datetime import datetime
import mammoth
from weasyprint import HTML

# Dati di accesso predefiniti
DEFAULT_USERNAME = "admin"
DEFAULT_PASSWORD = "password"

# Inizializza le variabili di sessione
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if 'username' not in st.session_state:
    st.session_state['username'] = DEFAULT_USERNAME

if 'password' not in st.session_state:
    st.session_state['password'] = DEFAULT_PASSWORD

# Funzione per sostituire NaN con stringa vuota
def valore_o_spazio(valore):
    return "" if pd.isna(valore) or valore == "nan" or valore is None else str(valore)

# Funzione per rimuovere i decimali
def rimuovi_decimali(valore):
    try:
        return str(int(float(valore)))
    except (ValueError, TypeError):
        return str(valore)

def formatta_pod(valore):
    try:
        if isinstance(valore, (int, float)):
            valore = '{:0.0f}'.format(float(valore))
        return str(valore).upper().strip()
    except (ValueError, TypeError):
        return str(valore)

# Funzione per formattare i numeri come stringa senza notazione scientifica
def formatta_numero_intero(numero):
    try:
        return '{:d}'.format(int(float(numero)))
    except ValueError:
        return str(numero)

# Funzione per formattare le date in italiano
def formatta_data_italiana(data):
    mesi = [
        "gennaio", "febbraio", "marzo", "aprile", "maggio", "giugno",
        "luglio", "agosto", "settembre", "ottobre", "novembre", "dicembre"
    ]
    giorno = data.day
    mese = mesi[data.month - 1]
    anno = data.year
    return f"{giorno} {mese} {anno}"

# Funzione per caricare i file Excel o CSV
def carica_file():
    uploaded_file1 = st.file_uploader("Carica il file ANAGRAFICHE", type=["xlsx", "csv"])
    uploaded_file2 = st.file_uploader("Carica il file FATTURE", type=["xlsx", "csv"])
    uploaded_file3 = st.file_uploader("Carica il file PRATICHE", type=["xlsx", "csv"])
    
    def leggi_file(file):
        if file is not None:
            try:
                if file.name.endswith('.xlsx'):
                    return pd.read_excel(file, na_values=["", "null", "nan", "NaN"], keep_default_na=False)
                elif file.name.endswith('.csv'):
                    return pd.read_csv(file, sep=';', na_values=["", "null", "nan", "NaN"], keep_default_na=False)
            except Exception as e:
                st.error(f"Errore nel caricamento del file {file.name}: {e}")
        return None

    df1 = leggi_file(uploaded_file1)
    df2 = leggi_file(uploaded_file2)
    df3 = leggi_file(uploaded_file3)
    
    return df1, df2, df3

# Funzione per normalizzare i nomi delle colonne
def normalizza_colonne(df):
    df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_').str.replace('.', '').str.replace("'", "")
    return df

# Funzione per compilare la tabella nel documento Word
def compila_tabella_esistente(doc, df_combinato):
    try:
        df_combinato = df_combinato.dropna(subset=['n_documento'])
        table = doc.tables[0]
        righe_da_aggiungere = len(df_combinato)
        righe_esistenti = len(table.rows)

        if righe_da_aggiungere > righe_esistenti - 1:
            for _ in range(righe_da_aggiungere - (righe_esistenti - 1)):
                table.add_row()

        for i, row in enumerate(df_combinato.itertuples(), 1):
            cells = table.rows[i].cells
            cells[0].text = str(row.data_reg)  
            cells[1].text = str(row.scadnetto)
            cells[2].text = str(row.n_documento)
            cells[3].text = str(row.importo_totale)
            cells[4].text = str(row.importo_pagato_totale)
            cells[5].text = str(row.residuo_ad_oggi)
            pod_value = row.pod
            cells[6].text = formatta_pod(pod_value)

    except Exception as e:
        st.error(f"Errore durante la compilazione della tabella: {e}")

# Funzione per generare il documento Word
def genera_documento_word(dati, df_combinato, template_path="decreto.docx"):
    doc = Document(template_path)
    data_generazione = formatta_data_italiana(datetime.now())

    placeholders = {
        "{ragione_sociale}": valore_o_spazio(dati.get('ragione_sociale_x','')),
        "{codice_fiscale}": rimuovi_decimali(valore_o_spazio(dati.get('codice_fiscale', ''))),
        "{partita_iva}": rimuovi_decimali(valore_o_spazio(dati.get('partita_iva', ''))),
        "{comune_residenza}": valore_o_spazio(dati.get('comune_residenza', '')),
        "{cap_residenza}": rimuovi_decimali(valore_o_spazio(dati.get('cap_residenza', 0))),
        "{indirizzo_residenza}": valore_o_spazio(dati.get('indirizzo_residenza', '')),
        "{settore_contabile}": valore_o_spazio(dati.get('settore_contabile', '')),
        "{codice_commerciale}": rimuovi_decimali(valore_o_spazio(dati.get('codice_commerciale', ''))),
        "{codice_soggetto}": rimuovi_decimali(valore_o_spazio(dati.get('codice_soggetto', 0))),
        "{comune_fornitura}": valore_o_spazio(dati.get('comune_fornitura', '')),
        "{provincia_fornitura}": valore_o_spazio(dati.get('provincia_fornitura', 'NA')).replace('nan', ''),
        "{indirizzo_fornitura}": valore_o_spazio(dati.get('indirizzo_fornitura', '')),
        "{data_generazione}": data_generazione,
        "{provincia_residenza}": valore_o_spazio(dati.get('provincia_residenza', 'NA')).replace('nan', ''),
        "{pod}": formatta_pod(valore_o_spazio(dati.get('pod', ''))),
        "{residuo_ad_oggi}": rimuovi_decimali(valore_o_spazio(dati.get('residuo_ad_oggi', ''))),
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    compila_tabella_esistente(doc, df_combinato)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Funzione per convertire Word in HTML con mammoth
def converti_word_in_html(file):
    with file:
        result = mammoth.convert_to_html(file)
        html = result.value  # Il contenuto HTML estratto
        return html

# Funzione per generare PDF da HTML con WeasyPrint
def genera_pdf_da_html(html_content):
    buffer = BytesIO()
    HTML(string=html_content).write_pdf(buffer)
    buffer.seek(0)
    return buffer

# Funzione di login
def login():
    st.title("Benvenuti al Portale Ricorsi D.I. dello Studio Carotenuto")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if username == st.session_state['username'] and password == st.session_state['password']:
            st.session_state['authenticated'] = True
            st.success("Accesso per Studio Carotenuto!")
        else:
            st.error("Username o password errati!")

# Funzione principale dell'applicazione
def main():
    st.markdown("[Clicca qui per collegarti al Portale Automazione Lettere Diffida](https://pagp-slcarotenuto.streamlit.app/)", unsafe_allow_html=True)

    if st.session_state['authenticated']:
        st.sidebar.title("Portale ricorso D.I. | Studio Carotenuto")

        st.title("Generatore automatico di Ricorsi D.I.")
        uploaded_word = st.file_uploader("Carica un documento Word (.docx)", type=["docx"])
        
        if uploaded_word is not None:
            st.write("File caricato con successo!")

            # Converte il file Word in HTML
            html_content = converti_word_in_html(uploaded_word)

            # Visualizza l'anteprima dell'HTML
            st.markdown(html_content, unsafe_allow_html=True)

            if st.button("Genera PDF"):
                pdf_buffer = genera_pdf_da_html(html_content)
                
                st.download_button(
                    label="Scarica il PDF generato",
                    data=pdf_buffer,
                    file_name="documento_generato.pdf",
                    mime="application/pdf"
                )
    else:
        login()

if __name__ == "__main__":
    main()
